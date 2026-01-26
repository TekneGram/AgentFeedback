# editing_class_new.py
from __future__ import annotations

import re
import difflib
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class TrackChangesEditor:
    """
    Emits Word track-changes markup (<w:ins>, <w:del>) into a NEW output .docx.

    Key points:
    - Revision IDs reset per output document (so multiple files in a run are safe).
    - Sentence splitting is regex-based (consistent with your original editing.py).
    - Can diff a whole doc as one paragraph OR paragraph-by-paragraph.
    - Can build a "single paragraph" report with ORIGINAL/EDITED/CORRECTED sections.
    """

    def __init__(self, author: str = "EssayLens", date: Optional[str] = None) -> None:
        self.author = author
        self.date = date or datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        self._rev_id = 1

    # ====
    # Sentence splitting (same behavior as your original)
    # ====
    _sentence_endings = re.compile(r"(?<=[.!?])\s+")

    @classmethod
    def split_into_sentences(cls, text: str) -> List[str]:
        txt = (text or "").strip()
        if not txt:
            return []
        return cls._sentence_endings.split(txt)

    # ============================================================
    # Track-changes helpers (proper WordprocessingML revisions)
    # ============================================================

    def reset_rev_ids(self) -> None:
        self._rev_id = 1

    def next_rev_id(self) -> int:
        rid = self._rev_id
        self._rev_id += 1
        return rid

    @staticmethod
    def enable_track_revisions(doc: Document) -> None:
        settings = doc.settings._element
        if settings.find(qn("w:trackRevisions")) is None:
            settings.append(OxmlElement("w:trackRevisions"))

    @staticmethod
    def append_plain_run(paragraph, text: str) -> None:
        if not text:
            return
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        paragraph._p.append(r)

    def add_tracked_insertion(self, paragraph, text: str) -> None:
        if not text:
            return
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(self.next_rev_id()))
        ins.set(qn("w:author"), self.author)
        ins.set(qn("w:date"), self.date)

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)

        ins.append(r)
        paragraph._p.append(ins)

    def add_tracked_deletion(self, paragraph, text: str) -> None:
        if not text or not text.strip():
            return

        delete = OxmlElement("w:del")
        delete.set(qn("w:id"), str(self.next_rev_id()))
        delete.set(qn("w:author"), self.author)
        delete.set(qn("w:date"), self.date)

        r = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = text
        r.append(del_text)

        delete.append(r)
        paragraph._p.append(delete)

    # ============================================================
    # Diff logic
    # ============================================================

    def apply_word_diff(self, paragraph, original: str, edited: str) -> None:
        """
        Emit [plain][del][ins] at token (word) level within a sentence pair.
        """
        orig_tokens = (original or "").split()
        edit_tokens = (edited or "").split()
        matcher = difflib.SequenceMatcher(None, orig_tokens, edit_tokens)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                self.append_plain_run(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
            elif tag == "delete":
                self.add_tracked_deletion(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
            elif tag == "insert":
                self.add_tracked_insertion(paragraph, " ".join(edit_tokens[j1:j2]) + " ")
            elif tag == "replace":
                self.add_tracked_deletion(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
                self.add_tracked_insertion(paragraph, " ".join(edit_tokens[j1:j2]) + " ")

    def apply_sentence_aligned_diff(self, paragraph, original_text: str, edited_text: str) -> None:
        """
        Align by sentence to keep diffs sane.
        - equal: output plain sentence
        - delete: output tracked deletion for whole sentence
        - insert: output tracked insertion for whole sentence
        - replace: pair sentences and do word-diff; leftovers become whole-sentence del/ins
        """
        original_sentences = self.split_into_sentences(original_text)
        edited_sentences = self.split_into_sentences(edited_text)

        sent_matcher = difflib.SequenceMatcher(None, original_sentences, edited_sentences)

        for tag, i1, i2, j1, j2 in sent_matcher.get_opcodes():
            if tag == "equal":
                for s in original_sentences[i1:i2]:
                    self.append_plain_run(paragraph, s + " ")

            elif tag == "delete":
                for s in original_sentences[i1:i2]:
                    self.add_tracked_deletion(paragraph, s + " ")

            elif tag == "insert":
                for s in edited_sentences[j1:j2]:
                    self.add_tracked_insertion(paragraph, s + " ")

            elif tag == "replace":
                pairs = min(i2 - i1, j2 - j1)

                # word-diff paired sentences
                for k in range(pairs):
                    self.apply_word_diff(paragraph, original_sentences[i1 + k], edited_sentences[j1 + k])

                # leftovers (whole-sentence changes)
                for s in original_sentences[i1 + pairs:i2]:
                    self.add_tracked_deletion(paragraph, s + " ")

                for s in edited_sentences[j1 + pairs:j2]:
                    self.add_tracked_insertion(paragraph, s + " ")

    # ============================================================
    # Builders
    # ============================================================

    def build_tracked_doc_flat(self, input_path: str, output_path: str, edited_text: str) -> None:
        self.reset_rev_ids()

        doc = Document(input_path)
        original_text = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

        out_doc = Document()
        self.enable_track_revisions(out_doc)

        p = out_doc.add_paragraph()
        self.apply_sentence_aligned_diff(p, original_text, (edited_text or "").strip())

        out_doc.save(output_path)

    def build_tracked_doc_from_paragraphs(
        self,
        input_path: str,
        output_path: str,
        edited_paragraphs: List[str],
        preserve_styles: bool = True,
        feedback_heading: str = "Cause–Effect Feedback",
        feedback_paragraphs: Optional[List[str]] = None,
        feedback_as_tracked_insertion: bool = True,
        add_page_break_before_feedback: bool = True,
    ) -> None:
        self.reset_rev_ids()

        orig_doc = Document(input_path)
        orig_paras = list(orig_doc.paragraphs)

        out_doc = Document()
        self.enable_track_revisions(out_doc)

        first_out_p = out_doc.paragraphs[0] if out_doc.paragraphs else out_doc.add_paragraph()
        used_first = False

        n = max(len(orig_paras), len(edited_paragraphs))

        for i in range(n):
            orig_text = orig_paras[i].text if i < len(orig_paras) else ""
            edited_text = edited_paragraphs[i] if i < len(edited_paragraphs) else ""

            if not used_first:
                out_p = first_out_p
                used_first = True
            else:
                out_p = out_doc.add_paragraph()

            if preserve_styles and i < len(orig_paras):
                try:
                    out_p.style = orig_paras[i].style
                except Exception:
                    pass

            if i >= len(orig_paras) and edited_text.strip():
                if feedback_as_tracked_insertion:
                    self.add_tracked_insertion(out_p, edited_text.strip() + " ")
                else:
                    out_p.add_run(edited_text.strip())
                continue

            if i >= len(edited_paragraphs) and orig_text.strip():
                self.add_tracked_deletion(out_p, orig_text.strip() + " ")
                continue

            self.apply_sentence_aligned_diff(out_p, orig_text, edited_text)

        if feedback_paragraphs:
            if add_page_break_before_feedback:
                out_doc.add_page_break()

            heading_p = out_doc.add_paragraph()
            try:
                heading_p.style = "Heading 1"
            except Exception:
                pass

            if feedback_as_tracked_insertion:
                self.add_tracked_insertion(heading_p, feedback_heading)
            else:
                heading_p.add_run(feedback_heading)

            for line in feedback_paragraphs:
                line = (line or "").rstrip()
                if not line.strip():
                    out_doc.add_paragraph("")
                    continue

                is_h2 = line.startswith("## ")
                text = line[3:].strip() if is_h2 else line.strip()

                p = out_doc.add_paragraph()
                if is_h2:
                    try:
                        p.style = "Heading 2"
                    except Exception:
                        pass

                if feedback_as_tracked_insertion:
                    self.add_tracked_insertion(p, text)
                else:
                    p.add_run(text)

        out_doc.save(output_path)

    # -------------------------------
    # Single-paragraph report
    # -------------------------------
    def build_single_paragraph_report(
        self,
        output_path: str,
        original_paragraphs: List[str],
        edited_text: str,
        corrected_text: str,
        feedback_heading: str = "Language Feedback",
        feedback_paragraphs: Optional[List[str]] = None,
        feedback_as_tracked_insertion: bool = False,
        add_page_break_before_feedback: bool = True,
        include_edited_text_section: bool = True,
    ) -> None:
        self.reset_rev_ids()

        out_doc = Document()
        self.enable_track_revisions(out_doc)

        def add_h1(title: str) -> None:
            p = out_doc.add_paragraph()
            try:
                p.style = "Heading 1"
            except Exception:
                pass
            p.add_run(title)

        # ORIGINAL TEXT (raw)
        add_h1("ORIGINAL TEXT")
        if original_paragraphs:
            for ptxt in original_paragraphs:
                out_doc.add_paragraph(ptxt or "")
        else:
            out_doc.add_paragraph("")

        out_doc.add_page_break()

        # EDITED TEXT (optional)
        if include_edited_text_section:
            add_h1("EDITED TEXT")
            out_doc.add_paragraph((edited_text or "").strip())
            out_doc.add_page_break()

        # CORRECTED TEXT (track changes vs edited)
        add_h1("CORRECTED TEXT")
        diff_p = out_doc.add_paragraph()
        self.apply_sentence_aligned_diff(
            diff_p,
            (edited_text or "").strip(),
            (corrected_text or "").strip(),
        )

        # Feedback section
        if feedback_paragraphs:
            if add_page_break_before_feedback:
                out_doc.add_page_break()

            add_h1(feedback_heading)

            for line in feedback_paragraphs:
                line = (line or "").rstrip()
                if not line.strip():
                    out_doc.add_paragraph("")
                    continue

                is_h2 = line.startswith("## ")
                text = line[3:].strip() if is_h2 else line.strip()

                p = out_doc.add_paragraph()
                if is_h2:
                    try:
                        p.style = "Heading 2"
                    except Exception:
                        pass

                if feedback_as_tracked_insertion:
                    self.add_tracked_insertion(p, text)
                else:
                    p.add_run(text)

        out_doc.save(output_path)


def build_tracked_doc(input_path: str, output_path: str, edited_text: str, author: str = "EssayLens") -> None:
    editor = TrackChangesEditor(author=author)
    editor.build_tracked_doc_flat(input_path, output_path, edited_text)
